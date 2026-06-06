"""Generate the Subprime Auto Credit Index project write-up as a Word document.

Reproducible build of ``Analysis/Subprime Auto Credit Index - Project Overview.docx``.
Content is grounded in the v1 implementation (universe.py / metrics.py /
index.py) and the design plan. Run: ``python Analysis/generate_writeup.py``.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Subprime Auto Credit Index - Project Overview.docx"

ACCENT = RGBColor(0x8B, 0x1A, 0x1A)  # deep crimson to match the index plot


def _style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _title(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = t.add_run("Subprime Auto Credit Index")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    s = sub.add_run("Project Overview — construction, mechanics, and a path to an "
                    "on-chain test market")
    s.italic = True
    s.font.size = Pt(12)
    meta = doc.add_paragraph()
    m = meta.add_run("Rivet  ·  Internal project memo  ·  June 2026")
    m.font.size = Pt(9)
    m.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()


def _h(doc: Document, n: int, text: str):
    h = doc.add_heading(text, level=n)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def _p(doc: Document, text: str):
    return doc.add_paragraph(text)


def _b(doc: Document, text: str):
    return doc.add_paragraph(text, style="List Bullet")


def _lead(doc: Document, lead: str, rest: str):
    """A bullet with a bold lead-in phrase."""
    para = doc.add_paragraph(style="List Bullet")
    r = para.add_run(lead)
    r.bold = True
    para.add_run(rest)
    return para


def build() -> None:
    doc = Document()
    _style(doc)
    _title(doc)

    # --- Executive summary -------------------------------------------------
    _h(doc, 1, "Executive Summary")
    _p(doc,
       "This project builds a monthly credit-deterioration index for U.S. subprime "
       "auto loans, computed directly from SEC asset-level (ABS-EE) disclosures. "
       "Unlike published delinquency statistics, which are aggregated and lagged, the "
       "index is derived from the loan-level cashflows of the actual securitized pools "
       "— several million individual subprime auto loans reported every month. The "
       "headline mark is a normalized stress signal that sits near zero in normal "
       "conditions and rises as the subprime consumer deteriorates. Because the index "
       "is a deterministic, fully reproducible function of public data, it is unusually "
       "well suited to on-chain settlement: any party can recompute and verify it, "
       "which is exactly the property a decentralized oracle requires. The remainder of "
       "this memo describes what the underlying data is, how the universe and the marks "
       "are constructed, how faithfully the index tracks subprime consumer credit, and "
       "how it could be traded via a smart contract — closing with the institutions "
       "most likely to take the other side of a first test trade.")

    # --- 1. ABS-EE files ---------------------------------------------------
    _h(doc, 1, "1.  ABS-EE Files — What They Are")
    _p(doc,
       "ABS-EE (“Asset-Backed Securities – Electronic Exhibits”) is the SEC "
       "form that carries asset-level disclosure for publicly registered "
       "securitizations, mandated by Regulation AB II. Since required compliance in "
       "late 2016, issuers of public ABS backed by auto loans, auto leases, "
       "residential and commercial mortgages, and debt securities must file, every "
       "reporting period, a standardized record for every individual asset in the pool.")
    _lead(doc, "Granularity. ",
          "The data lives in the EX-102 exhibit as XML conforming to the EDGAR ABS "
          "XML Technical Specification — one row per loan per month, with ~70+ "
          "standardized fields per row. For autos these span origination "
          "characteristics (obligor FICO, APR, original term, LTV, vehicle make/model/"
          "year, geography, income/employment verification) and monthly performance "
          "(scheduled vs. actual principal and interest, current delinquency status, "
          "charged-off principal, recoveries, repossession proceeds, modifications, and "
          "zero-balance / payoff codes).")
    _lead(doc, "Why it matters. ",
          "This is the rawest, most granular public credit data that exists. Most "
          "credit indicators are pool- or survey-aggregated; ABS-EE is the loan tape "
          "itself. It is timelier and finer than the New York Fed’s household-debt "
          "series, and — critically — it is public, so an index built on it "
          "carries no proprietary-data licensing constraints.")
    _lead(doc, "Timing. ",
          "Filings land roughly 15–25 days after each monthly reporting period, so "
          "the index for month M is computable in late month M+1. Subprime auto issuers "
          "filing ABS-EE include Santander, GM Financial / AmeriCredit, Exeter, "
          "DriveTime / Bridgecrest, and Carvana.")
    _lead(doc, "This project’s pipeline. ",
          "A downloader pulls EX-102 XML from EDGAR full-text search; parser modules "
          "clean and normalize the tape (reconciling balance walks, splitting consumer "
          "vs. commercial scores, deriving delinquency state, age, LTV, net losses, and "
          "prepayments); the reporting layer produces the per-loan / per-month frame the "
          "index consumes.")

    # --- 2. Index construction --------------------------------------------
    _h(doc, 1, "2.  Index Construction — What Loans Are Included")
    _p(doc,
       "The constituent universe is defined by a rule rather than a hand-picked list, "
       "so it self-updates as new subprime trusts begin filing and old ones amortize "
       "away. A securitization is included when it satisfies all of:")
    _lead(doc, "Asset class. ", "Auto loans only (leases are excluded from v1 because "
          "their residual-value risk is a different exposure from consumer credit).")
    _lead(doc, "Credit profile. ", "Weighted-average issuance consumer FICO below 640 "
          "— the conventional subprime line. (620 would isolate deep subprime; 660 "
          "would bleed into near-prime. 640 maximizes the subprime signal-to-noise.)")
    _lead(doc, "Size. ", "Original pool greater than $200 million, so small, "
          "idiosyncratic deals can’t dominate the signal.")
    _p(doc, "Current qualifying shelves include Santander Drive Auto Receivables Trust, "
            "AmeriCredit Automobile Receivables Trust (GM Financial subprime), Exeter "
            "Automobile Receivables Trust, Bridgecrest Lending Auto Securitization Trust "
            "(DriveTime), and the Carvana Auto Receivables Trust N-series. The set is "
            "derived from the data, not coded in.")
    _lead(doc, "Survivorship control. ",
          "Each securitized pool is static — it only shrinks. An aging deal looks "
          "worse purely from selection, because the good loans prepay and the weak ones "
          "remain. To prevent this from masquerading as market deterioration, a trust "
          "leaves the universe once its outstanding balance falls below 10% of its "
          "original size.")
    _lead(doc, "Pooling. ",
          "Deal boundaries are dissolved: the index is computed on the merged "
          "underlying loan universe, balance-weighted by current outstanding. Equal-"
          "weighting deals of different ages would inject seasoning bias; original-"
          "issuance weighting would freeze constituent influence at launch. Current-"
          "balance weighting is the right default for a deterioration signal.")

    # --- 3. Index mechanics ------------------------------------------------
    _h(doc, 1, "3.  Index Mechanics — How the Marks Are Calculated")
    _p(doc, "The index is built in two layers so each can be validated independently.")
    _h(doc, 2, "Performance layer (loan-level pooled)")
    _p(doc, "Five components are computed every month on the pooled loan universe, "
            "mixing leading and lagging signals so the index does not over-react to any "
            "one measure:")
    _lead(doc, "30+ and 60+ DPD balance share — ", "the standard stock measures of "
          "delinquency (leading / headline).")
    _lead(doc, "Current→30 monthly roll rate — ", "the share of currently-"
          "performing balance that rolls into 30-day delinquency the following month. "
          "This is the highest-value, most-leading component: it measures the flow into "
          "trouble, not the accumulated stock, and it turns before the level measures do.")
    _lead(doc, "Annualized net loss rate — ", "charge-offs net of recoveries, "
          "annualized on the beginning pool balance (coincident; the dollars actually "
          "lost).")
    _lead(doc, "Recovery rate — ", "recoveries divided by charge-offs, a read on "
          "loss severity / LGD direction.")
    _p(doc, "Pooling is exact rather than an average of ratios: for each component, the "
            "dollar numerator and denominator are summed across all live trusts in the "
            "month and then divided — algebraically identical to recomputing the "
            "metric on the combined loan pool.")
    _h(doc, 2, "Stress layer (the headline mark)")
    _p(doc, "The performance components run at structurally different levels (subprime "
            "auto loses 5–7% a year even in good times), so a raw level is not "
            "actionable. Each component is therefore converted to a rolling 24-month "
            "Z-score against its own trailing baseline (causal; a minimum 12-month "
            "window is required before any month is scored), clamped at ±3σ so a "
            "single outlier month cannot blow up the level, and oriented so that higher "
            "always means worse — the recovery-rate score is sign-flipped, since "
            "falling recoveries are deterioration. The component Z-scores are combined "
            "by inverse-volatility weighting (so a noisy metric doesn’t dominate) "
            "into a single composite. The result hovers near zero in normal periods and "
            "spikes positive when subprime auto is deteriorating — conceptually a "
            "VIX for subprime consumer credit.")
    _lead(doc, "COVID handling. ", "The April–December 2020 forbearance window is a "
          "structural break, not a market signal. It is flagged (and shaded on charts) "
          "but kept in the baseline, so it remains visible and can be excluded per "
          "analysis rather than silently dropped or interpolated.")
    _lead(doc, "Cadence and robustness. ", "Marks are monthly, published ~late M+1 with "
          "the filing lag. Servicers differ on when they charge off (90 vs. 120+ days); "
          "the roll-rate and delinquency components are insensitive to that policy "
          "choice, which is why they anchor the composite while the raw net-loss rate is "
          "treated with more caution.")

    # --- 4. Relationship to consumer credit --------------------------------
    _h(doc, 1, "4.  Index Relationship to Consumer Credit")
    _p(doc,
       "The index does not proxy subprime consumer credit — it is a direct "
       "measurement of it. Auto loans are an unusually clean window onto the subprime "
       "household: a car is typically the asset a subprime borrower most needs to keep "
       "(it secures access to work), so auto payment behavior tends to deteriorate "
       "early and visibly when household finances tighten, and it is reported monthly at "
       "the loan level.")
    _lead(doc, "Expected behavior. ", "A faithful index should reproduce the Q2-2020 "
          "delinquency spike and the subsequent stimulus / accommodation-driven dip, and "
          "the steady secular rise in subprime auto delinquency through 2023–2024 as "
          "pandemic savings were exhausted and payments normalized higher.")
    _lead(doc, "It should lead, not coincide. ", "Because the composite is anchored on "
          "the Current→30 roll rate — a flow measure — it should turn "
          "ahead of the stock-based delinquency levels that the New York Fed reports "
          "quarterly in its Household Debt and Credit Report. Loan-level monthly data is "
          "both finer and timelier than aggregate survey series.")
    _lead(doc, "Validation plan. ", "Benchmark qualitatively and statistically against "
          "(i) the NY Fed subprime-auto delinquency-flow series — expect "
          "correlation with the index leading by one to two quarters; (ii) ICE BofA "
          "high-yield auto OAS and KBRA’s quarterly subprime-auto trend reports — "
          "expect positive correlation without collapsing to equivalence (the index "
          "carries idiosyncratic, name-level information those spreads do not).")
    _lead(doc, "Known limitations. ", "ABS-EE history begins in 2016, so the baseline "
          "contains the 2020 shock but not the 2008–10 financial crisis; the "
          "reporting lag caps freshness at ~M+1; and the constituent set shifts as deals "
          "enter and amortize out, which the universe rule manages but does not "
          "eliminate.")

    # --- 5. DeFi / smart contracts ----------------------------------------
    _h(doc, 1, "5.  Affinity to DeFi and Smart Contracts")
    _p(doc,
       "The defining feature of this index for on-chain use is that it is a "
       "deterministic function of public inputs. The source data (EX-102 filings on "
       "EDGAR) and the calculation are both open, so any participant can independently "
       "recompute the mark and verify it bit-for-bit. That removes the trust assumption "
       "that normally forces credit derivatives into a bilateral, dealer-intermediated, "
       "ISDA-governed structure.")
    _lead(doc, "Oracle. ", "Publish the monthly composite (and its components) on-chain "
          "through a signed price feed or a Chainlink-style external adapter that runs "
          "the open-source calculator. Disputes can be resolved with an optimistic "
          "oracle (e.g. UMA): because the index is recomputable from public data, a "
          "challenge resolves deterministically against the reference implementation, "
          "rather than relying on a quorum’s opinion.")
    _lead(doc, "Instruments. ", "A cash-settled, monthly-settled future or funding-rate "
          "perpetual on the index level is the natural primitive — going long "
          "expresses a deterioration / credit-protection view, going short expresses "
          "stability or improvement. On top of that, range and binary options "
          "(“stress > 2σ by year-end”) and a collateralized synthetic token "
          "whose redemption tracks the index (minted against stablecoin collateral in a "
          "vault) extend the market to options and structured exposure.")
    _lead(doc, "Why monthly cadence fits. ", "The index moves slowly and updates "
          "monthly, so a monthly-settled future or a low-frequency funding perp is a "
          "better match than a high-frequency perpetual — settlement aligns with the "
          "filing cycle, and funding can be tied to the realized monthly mark.")
    _lead(doc, "Precedent. ", "This is the trustless analogue of the synthetic credit "
          "indices that traded subprime risk in TradFi — ABX.HE for subprime RMBS, "
          "CMBX for CRE — but settled on transparent public data and collateralized "
          "on-chain rather than intermediated by dealers. It lets a transparent, "
          "collateralized credit-risk-transfer market exist without the bilateral "
          "infrastructure that gates the cash ABS and CDS markets.")

    # --- 6. Targeted institutions -----------------------------------------
    _h(doc, 1, "6.  Targeted Institutions for a Test Trade")
    _p(doc, "A two-sided test market needs a natural hedger (buys protection / goes long "
            "deterioration) and a natural yield-seeker (sells protection / goes short).")
    _h(doc, 2, "Natural protection buyers (long the index)")
    _lead(doc, "Subprime auto lenders and issuers — ", "Santander Consumer, GM "
          "Financial, Exeter, DriveTime, Carvana. They carry warehouse, residual, and "
          "securitization-pipeline exposure to exactly this risk and currently have no "
          "clean, liquid macro hedge for it.")
    _lead(doc, "ABS investors in subordinate / residual tranches — ", "holders of "
          "the mezzanine and equity of subprime auto deals seeking to hedge mark-to-"
          "market without selling illiquid cash bonds.")
    _lead(doc, "Floorplan and dealer financiers — ", "lenders with concentrated "
          "subprime borrower exposure on the consumer side.")
    _h(doc, 2, "Natural protection sellers (short the index)")
    _lead(doc, "Credit hedge funds and RV desks — ", "a transparent, liquid way to "
          "express a constructive subprime-auto view or run relative value against cash "
          "ABS, without warehousing bonds.")
    _lead(doc, "DeFi-native credit and RWA protocols — ", "on-chain credit desks "
          "and real-world-asset platforms seeking verifiable real-world credit exposure "
          "with data they can audit.")
    _h(doc, 2, "First test-trade candidates")
    _p(doc, "The fastest-moving counterparties pair a crypto-forward risk appetite with "
            "genuine exposure. A practical seed pairing: one subprime auto lender or ABS "
            "residual holder with a hedging mandate on the long side, against a "
            "crypto-native credit fund or RWA protocol on the short side; settled "
            "monthly through an optimistic oracle, sized small, with the open-source "
            "calculator as the agreed reference. Oracle and settlement tooling (UMA, "
            "Chainlink) and the RWA-credit venues (Maple, Centrifuge, and similar) are "
            "the natural infrastructure partners for the pilot.")

    # --- Status note -------------------------------------------------------
    _h(doc, 1, "Current Status")
    _p(doc, "The v1 index engine — universe construction, per-trust monthly "
            "metrics, loan-level pooling, and the rolling-Z stress layer — is "
            "implemented and unit-tested. The next milestone is acquiring the full "
            "subprime constituent history from EDGAR and running the empirical "
            "validation against the NY Fed and HY-auto-spread benchmarks described in "
            "Section 4, after which a small on-chain test market can be scoped.")

    foot = doc.add_paragraph()
    f = foot.add_run("Internal working memo prepared for project planning. Not "
                     "investment advice and not an offer to transact.")
    f.italic = True
    f.font.size = Pt(8)
    f.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
